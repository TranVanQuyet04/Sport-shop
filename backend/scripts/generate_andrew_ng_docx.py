# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "andrew-ng-ai-reviewer-tieng-viet.docx"


def set_font(style, name="Times New Roman", size=None, bold=None, italic=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Trang ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


sections = [
    (
        "1. Andrew Ng là ai?",
        [
            "Andrew Ng, tên đầy đủ là Andrew Yan-Tak Ng, là một nhà khoa học máy tính, giảng viên, doanh nhân công nghệ và nhà truyền bá trí tuệ nhân tạo. Ông là một trong những nhân vật có ảnh hưởng lớn trong sự phát triển và phổ cập của machine learning và deep learning hiện đại.",
            "Andrew Ng được biết đến nhiều nhất qua các vai trò: đồng sáng lập và lãnh đạo dự án Google Brain, đồng sáng lập nền tảng học trực tuyến Coursera, sáng lập DeepLearning.AI, từng là Chief Scientist tại Baidu, sáng lập Landing AI, tham gia xây dựng AI Fund, và giảng dạy AI cho hàng triệu người học trên toàn cầu.",
            "Nói ngắn gọn, Andrew Ng là một trong những người giúp đưa AI từ phòng nghiên cứu ra đại chúng, từ các bài báo học thuật thành các khóa học, sản phẩm và ứng dụng thực tế.",
        ],
    ),
    (
        "2. Nền tảng học vấn và sự nghiệp ban đầu",
        [
            "Andrew Ng sinh năm 1976. Ông học tại những trường đại học rất mạnh về khoa học máy tính như Carnegie Mellon University, Massachusetts Institute of Technology (MIT) và University of California, Berkeley.",
            "Sau đó, ông trở thành giảng viên tại Stanford University, một trong những trung tâm nghiên cứu AI hàng đầu thế giới. Tại Stanford, Andrew Ng giảng dạy và nghiên cứu về machine learning, robotics và trí tuệ nhân tạo.",
            "Khóa học Machine Learning của ông tại Stanford và sau này trên Coursera đã trở thành một trong những khóa học nổi tiếng nhất trong lịch sử giáo dục trực tuyến.",
        ],
    ),
    (
        "3. Vai trò tại Stanford",
        [
            "Tại Stanford, Andrew Ng không chỉ là một nhà nghiên cứu mà còn là một người thầy có ảnh hưởng lớn. Ông nổi tiếng với khả năng giải thích các khái niệm khó trong machine learning một cách rõ ràng và có hệ thống.",
            "Những chủ đề thường gắn với các bài giảng của ông gồm linear regression, logistic regression, gradient descent, bias và variance, regularization, neural networks, support vector machines, unsupervised learning và recommender systems.",
            "Điều làm các khóa học của Andrew Ng nổi bật là cách ông giúp người học hiểu bản chất trước khi đi sâu vào công cụ. Vì vậy, nhiều người mới bắt đầu AI thường được khuyên học Andrew Ng đầu tiên.",
        ],
    ),
    (
        "4. Google Brain",
        [
            "Một dấu mốc quan trọng trong sự nghiệp của Andrew Ng là Google Brain. Đây là một dự án nghiên cứu deep learning tại Google, được xây dựng trong giai đoạn deep learning bắt đầu bùng nổ.",
            "Andrew Ng cùng các nhà nghiên cứu khác đã tham gia thúc đẩy việc huấn luyện các mạng neural quy mô lớn bằng dữ liệu lớn và năng lực tính toán lớn. Google Brain góp phần làm thay đổi cách nhìn của công nghiệp về deep learning.",
            "Vai trò của Andrew Ng ở Google Brain cho thấy ông không chỉ mạnh về giảng dạy, mà còn có đóng góp trong việc đưa deep learning vào môi trường công nghiệp quy mô lớn.",
        ],
    ),
    (
        "5. Coursera và giáo dục trực tuyến",
        [
            "Năm 2012, Andrew Ng đồng sáng lập Coursera cùng Daphne Koller. Coursera là một nền tảng học trực tuyến cho phép người học trên toàn thế giới tiếp cận các khóa học từ nhiều trường đại học và tổ chức lớn.",
            "Khóa Machine Learning của Andrew Ng trên Coursera trở thành một hiện tượng. Nó giúp rất nhiều sinh viên, lập trình viên, kỹ sư và người chuyển ngành tiếp cận AI một cách có hệ thống.",
            "Tác động của Coursera và các khóa học của Andrew Ng nằm ở việc làm AI trở nên dễ tiếp cận hơn, giảm rào cản học tập cho người không ở các trường hàng đầu, và tạo nên một con đường học machine learning phổ biến cho người mới.",
        ],
    ),
    (
        "6. Baidu",
        [
            "Sau Google, Andrew Ng từng làm Chief Scientist tại Baidu, một trong những công ty công nghệ lớn của Trung Quốc. Tại đây, ông tham gia xây dựng và lãnh đạo các nhóm nghiên cứu AI.",
            "Những lĩnh vực AI liên quan đến giai đoạn này bao gồm nhận dạng giọng nói, xử lý ngôn ngữ tự nhiên, thị giác máy tính, deep learning và các nền tảng AI ứng dụng.",
            "Giai đoạn Baidu cho thấy Andrew Ng có khả năng đưa AI vào các bài toán công nghiệp lớn, không chỉ dừng lại ở giáo dục hay nghiên cứu học thuật.",
        ],
    ),
    (
        "7. DeepLearning.AI",
        [
            "Sau khi rời Baidu, Andrew Ng sáng lập DeepLearning.AI. Đây là một tổ chức giáo dục tập trung vào AI, machine learning, deep learning và gần đây là generative AI.",
            "Một số khóa học nổi bật của DeepLearning.AI gồm Machine Learning Specialization, Deep Learning Specialization, AI for Everyone, Generative AI for Everyone, Machine Learning Engineering for Production, cùng các khóa ngắn về prompt engineering, large language models, retrieval augmented generation, evaluation và MLOps.",
            "DeepLearning.AI tiếp tục thể hiện mục tiêu lớn của Andrew Ng: dân chủ hóa giáo dục AI, giúp nhiều người hơn có thể hiểu và ứng dụng AI.",
        ],
    ),
    (
        "8. Landing AI và data-centric AI",
        [
            "Andrew Ng cũng là người sáng lập Landing AI, một công ty tập trung vào việc đưa AI vào doanh nghiệp, đặc biệt trong lĩnh vực thị giác máy tính và sản xuất.",
            "Một khái niệm rất gắn với Andrew Ng trong giai đoạn này là data-centric AI. Data-centric AI nhấn mạnh rằng để cải thiện hệ thống AI, không nên chỉ tập trung vào việc thay đổi model. Trong nhiều trường hợp, việc cải thiện dữ liệu còn quan trọng hơn.",
            "Dữ liệu cần sạch hơn, nhãn dữ liệu cần nhất quán hơn, quy trình thu thập dữ liệu cần tốt hơn, và lỗi của model cần được phân tích để cải thiện dữ liệu. Tư duy này rất thực tế trong môi trường doanh nghiệp, nơi dữ liệu thường không gọn gàng như các bộ dữ liệu mẫu trong sách giáo khoa.",
        ],
    ),
    (
        "9. AI Fund",
        [
            "AI Fund là một tổ chức đầu tư và xây dựng startup AI có liên quan đến Andrew Ng. Mục tiêu của AI Fund là tìm các cơ hội ứng dụng AI vào những bài toán thực tế và xây dựng công ty từ các cơ hội đó.",
            "Điều này phản ánh cách Andrew Ng nhìn về AI: AI không chỉ là một ngành nghiên cứu, mà là một nền tảng có thể tạo ra sản phẩm, doanh nghiệp và giá trị kinh tế mới.",
        ],
    ),
    (
        "10. Andrew Ng có phải “AI reviewer” không?",
        [
            "Cụm từ “Andrew Ng AI reviewer” có thể được hiểu theo nhiều cách.",
            "Nếu “reviewer” nghĩa là người phản biện trong học thuật, Andrew Ng có liên quan đến việc đánh giá và phản biện các công trình nghiên cứu AI trong sự nghiệp học thuật của mình. Ông từng có vai trò trong cộng đồng nghiên cứu, tạp chí và hội nghị liên quan đến AI, machine learning và robotics.",
            "Nếu “AI reviewer” nghĩa là một sản phẩm, công cụ hoặc ứng dụng có tên như vậy, hiện không có bằng chứng rõ ràng rằng “Andrew Ng AI Reviewer” là một sản phẩm chính thức nổi tiếng của Andrew Ng.",
            "Vì vậy, khi nhắc đến “Andrew Ng AI reviewer”, nên phân biệt giữa Andrew Ng như một chuyên gia AI có khả năng đánh giá các vấn đề AI, và “AI Reviewer” như một tên sản phẩm chưa được xác nhận rõ ràng.",
        ],
    ),
    (
        "11. Tư duy AI nổi bật của Andrew Ng",
        [
            "Andrew Ng từng so sánh AI với điện. Ý tưởng là AI có thể trở thành một công nghệ nền tảng, ảnh hưởng đến nhiều ngành khác nhau, giống như cách điện đã thay đổi công nghiệp, giao thông, sản xuất và đời sống.",
            "Ông nhấn mạnh rằng nhiều dự án AI thất bại không phải vì thiếu thuật toán phức tạp, mà vì dữ liệu kém chất lượng, bài toán không rõ, hoặc quy trình đánh giá không đúng.",
            "Theo tư duy của ông, một dự án AI nên bắt đầu từ các câu hỏi: bài toán cần giải quyết là gì, có dữ liệu phù hợp không, kết quả được đo lường bằng cách nào, và AI có thật sự tạo giá trị hơn cách làm hiện tại không.",
            "Andrew Ng cũng tin rằng AI không chỉ dành cho nhà nghiên cứu. Nhà quản lý, nhà sáng lập, kỹ sư, sinh viên và người làm trong nhiều lĩnh vực đều nên hiểu AI ở mức phù hợp với công việc của mình.",
        ],
    ),
    (
        "12. Điểm mạnh của Andrew Ng",
        [
            "Andrew Ng có nhiều điểm mạnh nổi bật: giải thích kiến thức khó một cách dễ hiểu, có nền tảng nghiên cứu vững chắc, có kinh nghiệm xây dựng AI trong công ty lớn, có khả năng biến kiến thức AI thành chương trình học có hệ thống, có tư duy thực dụng về ứng dụng AI trong doanh nghiệp, và có ảnh hưởng lớn đến cộng đồng người học AI toàn cầu.",
            "Nói ngắn gọn, Andrew Ng là cầu nối giữa nghiên cứu, giáo dục và ứng dụng AI.",
        ],
    ),
    (
        "13. Những điểm có thể bị phê bình",
        [
            "Cần nhìn Andrew Ng một cách cân bằng. Một số ý kiến cho rằng một số khóa học của ông có tính nền tảng, nên có thể chậm với người muốn đi thẳng vào large language models hoặc generative AI.",
            "Các bài tập trong khóa học thường được thiết kế gọn gàng, khác với môi trường production thật. Ngoài ra, ông có xu hướng lạc quan và thực dụng về AI ứng dụng, trong khi một số nhà nghiên cứu khác tập trung nhiều hơn vào rủi ro dài hạn của AI.",
            "Tuy nhiên, các điểm này không làm giảm giá trị của Andrew Ng. Nó chỉ cho thấy ông đại diện cho một góc nhìn cụ thể: học nền tảng, xây ứng dụng, cải thiện dữ liệu và đưa AI vào thực tế.",
        ],
    ),
    (
        "14. So sánh với một số nhân vật AI khác",
        [
            "Geoffrey Hinton nổi bật về nền tảng neural networks và deep learning. Yann LeCun nổi bật về computer vision, convolutional neural networks và AI tại Meta. Yoshua Bengio nổi bật về deep learning, representation learning và gần đây quan tâm nhiều đến AI safety.",
            "Fei-Fei Li nổi bật về computer vision, ImageNet và human-centered AI. Demis Hassabis nổi bật với DeepMind, AlphaGo và AlphaFold. Sam Altman nổi bật về OpenAI, ChatGPT và chiến lược sản phẩm generative AI.",
            "Trong bản đồ AI, Andrew Ng nổi bật về giáo dục AI, Google Brain, Coursera, DeepLearning.AI, data-centric AI và ứng dụng AI trong doanh nghiệp.",
        ],
    ),
    (
        "15. Nên học gì từ Andrew Ng?",
        [
            "Từ Andrew Ng, có thể rút ra một số bài học quan trọng: nên học nền tảng machine learning trước khi chạy theo công cụ mới; AI cần được gắn với bài toán thực tế; dữ liệu tốt có thể quan trọng ngang hoặc hơn model tốt; người không chuyên kỹ thuật vẫn nên hiểu AI ở mức chiến lược; học AI hiệu quả nhất khi kết hợp lý thuyết với project thực hành; và không nên xem AI là phép màu, mà cần có quy trình đánh giá, thử nghiệm và cải tiến.",
        ],
    ),
    (
        "16. Lộ trình tìm hiểu Andrew Ng và AI",
        [
            "Nếu mới bắt đầu tìm hiểu về Andrew Ng và AI, có thể đi theo thứ tự sau: xem các bài phỏng vấn Andrew Ng trên YouTube; học khóa AI for Everyone nếu chưa có nền tảng kỹ thuật; học Machine Learning Specialization nếu muốn bắt đầu nghiêm túc với machine learning; học Deep Learning Specialization nếu muốn hiểu neural networks; đọc Machine Learning Yearning; tìm hiểu về Google Brain; tìm hiểu DeepLearning.AI; và tìm hiểu data-centric AI cùng Landing AI.",
        ],
    ),
    (
        "17. Đánh giá tổng quan",
        [
            "Andrew Ng là một nhân vật rất đáng tìm hiểu nếu muốn hiểu AI hiện đại. Ông không phải người phát minh ra AI, và cũng không phải là người duy nhất tạo nên deep learning. Tuy nhiên, ông có vai trò rất lớn trong việc đưa AI đến với nhiều người hơn.",
            "Ảnh hưởng của Andrew Ng nằm ở ba trục lớn: nghiên cứu và công nghiệp, giáo dục, và tư duy ứng dụng. Ông gắn với Google Brain, Baidu, Landing AI, Coursera, DeepLearning.AI, các khóa học machine learning và deep learning, data-centric AI, AI cho doanh nghiệp, và cách nhìn AI như một công nghệ nền tảng.",
            "Tóm lại, Andrew Ng là một trong những người quan trọng nhất trong việc biến AI từ lĩnh vực khó tiếp cận thành một kỹ năng và công nghệ mà hàng triệu người có thể học, hiểu và ứng dụng.",
        ],
    ),
    (
        "18. Nguồn tham khảo nên đọc tiếp",
        [
            "Andrew Ng official website: https://www.andrewng.org/",
            "DeepLearning.AI: https://www.deeplearning.ai/",
            "Coursera instructor profile: https://www.coursera.org/instructor/andrewng",
            "Machine Learning Yearning: https://www.deeplearning.ai/machine-learning-yearning/",
            "The Batch newsletter: https://www.deeplearning.ai/the-batch/",
            "TIME100 AI profile: https://time.com/collection/time100-ai/",
        ],
    ),
]


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    set_font(doc.styles["Normal"], size=12)
    set_font(doc.styles["Title"], size=20, bold=True)
    set_font(doc.styles["Heading 1"], size=15, bold=True)
    set_font(doc.styles["Heading 2"], size=13, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Andrew Ng và vai trò trong lĩnh vực AI")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Tổng quan về tiểu sử, đóng góp, tư tưởng và ảnh hưởng")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Tài liệu tổng hợp về Andrew Ng và vai trò trong lĩnh vực AI"
    add_page_number(section.footer.add_paragraph())

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
